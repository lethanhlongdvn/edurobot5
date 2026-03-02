# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_edurobot_chuyen_de():
    doc = Document()
    
    # --- PHẦN TIÊU ĐỀ ---
    title = doc.add_heading(u'CHUYÊN ĐỀ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    main_title = doc.add_heading(u'ỨNG DỤNG NỀN TẢNG EDUROBOT-5 TRONG DẠY HỌC TOÁN LỚP 5 THEO HƯỚNG TƯƠNG TÁC VÀ PHÁT TRIỂN NĂNG LỰC SỐ', 1)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- PHẦN 1: SỰ CẦN THIẾT ---
    doc.add_heading(u'I. SỰ CẦN THIẾT CỦA CHUYÊN ĐỀ', level=1)
    doc.add_paragraph(u'Trong chương trình Giáo dục phổ thông 2018, môn Toán lớp 5 đóng vai trò then chốt trong việc hình thành tư duy logic và năng lực giải quyết vấn đề. Tuy nhiên, các phương pháp dạy học truyền thống bằng PowerPoint hiện nay vẫn còn mang tính "trình chiếu một chiều", học sinh chủ yếu là người quan sát, chưa thực sự được tham gia vào quá trình khám phá kiến thức.')
    doc.add_paragraph(u'EduRobot-5 ra đời như một giải pháp đột phá, chuyển đổi các bài giảng tĩnh thành không gian học tập số tương tác 100%. Vì vậy, việc triển khai chuyên đề ứng dụng EduRobot-5 là bước đi tất yếu để nâng cao chất lượng dạy học theo hướng hiện đại.')
    
    # --- PHẦN 2: MỤC TIÊU ---
    doc.add_heading(u'II. MỤC TIÊU CHUYÊN ĐỀ', level=1)
    p2 = doc.add_paragraph()
    p2.add_run(u'• Về kiến thức: ').bold = True
    p2.add_run(u'Số hóa 100% học liệu Toán 5 khớp với Phân phối chương trình (PPCT).')
    p2 = doc.add_paragraph()
    p2.add_run(u'• Về kỹ năng tương tác: ').bold = True
    p2.add_run(u'Học sinh được trực tiếp thao tác trên các mô hình hình học số và công thức kéo-thả.')
    p2 = doc.add_paragraph()
    p2.add_run(u'• Về đánh giá: ').bold = True
    p2.add_run(u'Sử dụng dữ liệu thời gian thực từ Quiz để điều chỉnh phương pháp dạy ngay trong tiết học.')
    
    # --- PHẦN 3: CÁC BIỆN PHÁP THỰC HIỆN VỚI EDUROBOT-5 ---
    doc.add_heading(u'III. CÁC BIỆN PHÁP THỰC HIỆN TRÊN NỀN TẢNG EDUROBOT-5', level=1)
    
    # BP1
    doc.add_heading(u'1. Biện pháp 1: Số hóa bài giảng với "Robot dẫn đường" (Interactive Lesson Engine)', level=2)
    doc.add_paragraph(u'Sử dụng Module bài học của EduRobot-5 để cấu trúc bài giảng theo 4 hoạt động của Bộ Giáo dục. Robot sẽ đóng vai trò người dẫn dắt, đưa ra các câu hỏi gợi mở qua Voice và Animation, giúp HS không bị nhàm chán.')
    
    # BP2
    doc.add_heading(u'2. Biện pháp 2: Tương tác trực quan với "Công cụ toán học số"', level=2)
    doc.add_paragraph(u'Đặc biệt hiệu quả cho các bài toán về Hình hộp chữ nhật, Hình lập phương (Tiết 116-120). Học sinh có thể tương tác với các hình khai triển (SVG/3D), thực hiện kéo- thả các kích thước vào công thức để tính diện tích/thể tích ngay trên màn hình.')
    
    # BP3
    doc.add_heading(u'3. Biện pháp 3: Gamification qua hệ thống Quiz thông minh', level=2)
    doc.add_paragraph(u'Hệ thống Quiz của EduRobot-5 cung cấp phản hồi tức thì (Instant Feedback). Robot sẽ chúc mừng khi HS chọn đúng hoặc đưa ra gợi ý khi HS chọn sai. Bảng xếp hạng trực tuyến giúp tạo tinh thần thi đua lành mạnh.')
    
    # BP4
    doc.add_heading(u'4. Biện pháp 4: Xây dựng kho học liệu số dùng chung và tự học', level=2)
    doc.add_paragraph(u'Nền tảng web-based cho phép giáo viên chia sẻ mã bài học cho đồng nghiệp và học sinh. HS có thể tự ôn tập (Blended Learning) tại nhà trên máy tính hoặc điện thoại.')
    
    # --- PHẦN 4: QUY TRÌNH DẠY HỌC ---
    doc.add_heading(u'IV. QUY TRÌNH TỔ CHỨC TIẾT DẠY TOÁN TƯƠNG TÁC', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'Hoạt động'
    hdr_cells[1].text = u'Ứng dụng trên EduRobot-5'
    
    activities = [
        (u'Khởi động', u'Robot đưa ra trò chơi nhanh (Mini-Quiz) để ôn lại kiến thức cũ.'),
        (u'Khám phá', u'HS quan sát mô hình động, trực tiếp thao tác để tự rút ra công thức.'),
        (u'Luyện tập', u'Làm bài tập tương tác kéo-thả, điền số. Robot chấm điểm ngay.'),
        (u'Vận dụng', u'HS giải quyết tình huống thực tế qua hình ảnh bài toán thực tế.'),
    ]
    
    for act, app in activities:
        row_cells = table.add_row().cells
        row_cells[0].text = act
        row_cells[1].text = app
        
    # --- PHẦN 5: KẾT LUẬN ---
    doc.add_heading(u'V. KẾT LUẬN', level=1)
    doc.add_paragraph(u'Việc ứng dụng EduRobot-5 trong dạy Toán lớp 5 không chỉ giúp đổi mới phương pháp giảng dạy mà còn tạo ra một hệ sinh thái học tập hiện đại. Đây là minh chứng cho việc làm chủ công nghệ của giáo viên, hướng tới mục tiêu phát triển năng lực số toàn diện cho học sinh Tiểu học.')
    
    # Lưu file
    file_name = 'CHUYEN_DE_TOAN_5_EDUROBOT_5.docx'
    doc.save(file_name)
    print(f'Created: {os.path.abspath(file_name)}')

if __name__ == "__main__":
    create_edurobot_chuyen_de()
