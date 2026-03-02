# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_plan_docx():
    doc = Document()
    
    # Thiết kế tiêu đề chính
    title = doc.add_heading(u'KẾ HOẠCH NÂNG CẤP DỰ ÁN EDUROBOT5', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Mục tiêu chung
    doc.add_heading(u'I. MỤC ĐÍCH VÀ Ý TƯỞNG CHỦ ĐẠO', level=1)
    p = doc.add_paragraph(u'Nâng cấp dự án EduRobot5 trở thành nền tảng giáo dục thông minh (Interactive Next-Gen Platform) phục vụ dạy học môn Toán lớp 5. Dự án tập trung giải quyết các vấn đề "thụ động" và "thiếu tính tương tác" đã được chỉ ra trong chuyên đề sinh hoạt chuyên môn.')
    
    # Các trụ cột nâng cấp
    doc.add_heading(u'II. CÁC TRỤ CỘT NÂNG CẤP CHÍNH', level=1)
    
    # 1. Gamification
    doc.add_heading(u'1. Nâng cấp linh hồn - Gamification & Phản hồi nhanh', level=2)
    p1 = doc.add_paragraph()
    p1.add_run(u'• Tính năng Phản hồi tức thì (Instant Feedback): ').bold = True
    p1.add_run(u'Tích hợp âm thanh (SFX) và Robot động (Lottie Animation) chúc mừng khi học sinh trả lời đúng.')
    p1 = doc.add_paragraph()
    p1.add_run(u'• Tính năng Streak & Điểm thưởng: ').bold = True
    p1.add_run(u'Tạo động lực khi học sinh trả lời đúng liên tiếp.')
    p1 = doc.add_paragraph()
    p1.add_run(u'• Bảng vàng xếp hạng (Leaderboard): ').bold = True
    p1.add_run(u'Hiển thị xếp hạng 10 bạn xuất sắc nhất lớp thời gian thực qua Firebase.')
    
    # 2. Interactive Tool
    doc.add_heading(u'2. Bộ công cụ tương tác chuyên nghiệp', level=2)
    p2 = doc.add_paragraph()
    p2.add_run(u'• Module Công thức Kéo-Thả (Drag & Drop): ').bold = True
    p2.add_run(u'Thay thế trắc nghiệm tĩnh bằng thao tác kéo số vào đúng vị trí công thức.')
    p2 = doc.add_paragraph()
    p2.add_run(u'• Phòng thí nghiệm hình học động: ').bold = True
    p2.add_run(u'HS thao tác thanh trượt để thay đổi bán kính/cạnh và thấy diện tích biến thiên thực tế (giả lập GeoGebra).')
    p2 = doc.add_paragraph()
    p2.add_run(u'• Dạng bài Ghép đôi (Matching): ').bold = True
    p2.add_run(u'Tương tác nối khái niệm với hình vẽ minh họa.')
    
    # 3. Visual & Media
    doc.add_heading(u'3. Thiết kế cao cấp & Học liệu thực tế', level=2)
    p3 = doc.add_paragraph()
    p3.add_run(u'• Design System (Glassmorphism): ').bold = True
    p3.add_run(u'Giao diện kính mờ, chuyển động mượt mà tạo cảm giác "vị lai".')
    p3 = doc.add_paragraph()
    p3.add_run(u'• Vận dụng thực tế: ').bold = True
    p3.add_run(u'HS có thể tải ảnh chụp bồn hoa, vật thể thực tế lên và đo đạc/tính toán trực tiếp trên ứng dụng.')
    
    # Lộ trình
    doc.add_heading(u'III. LỘ TRÌNH TRIỂN KHAI', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'Giai đoạn'
    hdr_cells[1].text = u'Nội dung trọng tâm'
    
    row_cells = table.add_row().cells
    row_cells[0].text = u'Tháng 1'
    row_cells[1].text = u'Hoàn thiện Quiz Engine 2.0 & Gamification'
    
    row_cells = table.add_row().cells
    row_cells[0].text = u'Tháng 2'
    row_cells[1].text = u'Phát triển Module Tương tác Drag-Drop & Hình học động'
    
    row_cells = table.add_row().cells
    row_cells[0].text = u'Tháng 3'
    row_cells[1].text = u'Đổi mới Giao diện cao cấp & Tích hợp học liệu số thực tế'
    
    # Lời kết
    doc.add_heading(u'IV. KẾT LUẬN', level=1)
    doc.add_paragraph(u'EduRobot5 không chỉ là một phần mềm học tập, mà là giải pháp chuyển đổi phương thức dạy toán lớp 5 từ thụ động sang chủ động khám phá, góp phần hiện thực hóa mục tiêu "Mỗi giờ học Toán là một giờ học đầy hứng thú".')
    
    # Lưu
    file_name = 'KE_HOACH_NANG_CAP_EDUROBOT5.docx'
    doc.save(file_name)
    print(f'Created: {os.path.abspath(file_name)}')

if __name__ == "__main__":
    create_plan_docx()
