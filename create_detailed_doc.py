# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import os

def set_font(run, font_name='Times New Roman', size=13):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)

def add_detailed_chuyen_de():
    doc = Document()
    
    # Thiết lập lề trang
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.75)

    # --- TRANG BÌA / TIÊU ĐỀ ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(u'PHÒNG GIÁO DỤC VÀ ĐÀO TẠO HUYỆN CÀNG LONG\nTRƯỜNG TIỂU HỌC ĐỖ VĂN NẠI\n----------***----------\n\n\n')
    set_font(run, size=14)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(u'BÁO CÁO CHUYÊN ĐỀ\n')
    set_font(run, size=18)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(u'ỨNG DỤNG NỀN TẢNG EDUROBOT-5 TRONG DẠY HỌC TOÁN LỚP 5\nTHEO HƯỚNG TƯƠNG TÁC VÀ PHÁT TRIỂN NĂNG LỰC SỐ\n')
    set_font(run, size=20)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(u'\n\n\nNgười thực hiện: Nhóm dự án EduRobot-5\nĐối tượng: Học sinh khối lớp 5\nThời gian: Năm học 2025 - 2026\n')
    set_font(run, size=13)
    
    doc.add_page_break()

    # --- PHẦN I: MỞ ĐẦU ---
    doc.add_heading(u'I. SỰ CẦN THIẾT VÀ LÝ DO CHỌN CHUYÊN ĐỀ', level=1)
    
    doc.add_heading(u'1. Tầm quan trọng của môn Toán lớp 5', level=2)
    p = doc.add_paragraph(u'Môn Toán ở bậc Tiểu học, đặc biệt là lớp 5, giữ vai trò "cửa ngõ" quan trọng giúp học sinh hình thành tư duy logic và khả năng giải quyết vấn đề thực tế. Các nội dung về số thập phân, hình học (diện tích, thể tích) và chuyển động đều là những kiến thức trừu tượng, đòi hỏi học sinh phải có khả năng tưởng tượng phong phú.')

    doc.add_heading(u'2. Thách thức trong dạy học hiện nay', level=2)
    p = doc.add_paragraph(u'Mặc dù giáo viên đã có ý thức ứng dụng CNTT như dùng máy chiếu hay PowerPoint, nhưng phần lớn vẫn chỉ dừng lại ở việc "Chiếu ảnh thay bảng đen". Học sinh vẫn ngồi nghe, nhìn hình ảnh tĩnh, dẫn đến tình trạng "biết lý thuyết nhưng khó thực hành". Yếu tố "Tương tác" – nghĩa là học sinh được chạm vào kiến thức, được thao tác và nhận được phản hồi ngay lập tức – vẫn còn rất thiếu.')

    doc.add_heading(u'3. Lý do lựa chọn giải pháp EduRobot-5', level=2)
    p = doc.add_paragraph(u'EduRobot-5 không đơn thuần là một trang web xem bài học, mà là một "Trợ lý giảng dạy số". Giải pháp này được chọn vì nó giải quyết đúng 3 nút thắt:\n'
                          u'- HS được trực tiếp thao tác (Kéo, thả, xoay hình).\n'
                          u'- GV không cần giỏi CNTT vẫn dùng được (Giao diện bằng tiếng Việt, Robot hướng dẫn tận nơi).\n'
                          u'- Đánh giá tự động (HS làm bài đến đâu, Robot sửa đến đó).')

    # --- PHẦN II: THỰC TRẠNG ---
    doc.add_heading(u'II. THỰC TRẠNG TÌNH HÌNH TẠI ĐƠN VỊ', level=1)
    
    doc.add_heading(u'1. Thuận lợi', level=2)
    p = doc.add_paragraph(u'- Cơ sở vật chất: Nhà trường đã được trang bị tivi thông minh, máy chiếu và đường truyền Internet.\n'
                          u'- Học sinh: Đa số các em đều đã quen thuộc với điện thoại, máy tính bảng nên rất nhanh nhạy với các nút bấm trên màn hình.\n'
                          u'- Phụ huynh: Rất đồng thuận trong việc cho con tiếp cận công nghệ để học tập thay vì giải trí vô bổ.')

    doc.add_heading(u'2. Khó khăn (Nguyên nhân cần thay đổi)', level=2)
    p = doc.add_paragraph(u'2.1. Về phía Giáo viên:\n'
                          u'- Trình độ công nghệ không đồng đều. Nhiều GV lớn tuổi e ngại các phần mềm phức tạp.\n'
                          u'- Việc soạn một bài giảng điện tử có sự tương tác (như làm trò chơi, làm hình ảnh động) tốn quá nhiều thời gian.\n\n'
                          u'2.2. Về phía Học sinh:\n'
                          u'- Khả năng tư duy trừu tượng chưa cao. Ví dụ: Khi học về hình hộp chữ nhật, các em khó hình dung các mặt khuất nếu chỉ nhìn hình vẽ trên bìa sách.')

    # --- PHẦN III: GIỚI THIỆU HỆ SINH THÁI EDUROBOT-5 ---
    doc.add_heading(u'III. NỀN TẢNG EDUROBOT-5 – ĐỘNG LỰC MỚI CHO TIẾT TOÁN', level=1)
    p = doc.add_paragraph(u'Để giáo viên hiểu rõ, EduRobot-5 là một hệ thống dạy học trực tuyến được thiết kế theo mô hình "Robot đồng hành". Nó bao gồm các tính năng chính:')
    
    doc.add_bullet_point(u'Robot hướng dẫn: Một nhân vật ảo xuyên suốt bài học, nhắc nhở nhiệm vụ bằng tiếng Việt.')
    doc.add_bullet_point(u'Lesson Engine (Bộ máy bài giảng): Chỉ cần một nút bấm "Tiếp tục", GV có thể trình diễn bài học theo đúng trình tự sư phạm.')
    doc.add_bullet_point(u'Kho dữ liệu số: Chứa sẵn hàng ngàn hình ảnh mẫu, video thực tế về toán học đã được cắt gọt từ SGK.')

    # --- PHẦN IV: CÁC BIỆN PHÁP THỰC HIỆN CHI TIẾT ---
    doc.add_heading(u'IV. CÁC BIỆN PHÁP TỔ CHỨC DẠY HỌC TƯƠNG TÁC TẬP TRUNG', level=1)

    # Biện pháp 1
    doc.add_heading(u'BIỆN PHÁP 1: Chuyển đổi bài giảng "Tĩnh" sang "Động" trực quan', level=2)
    p = doc.add_paragraph(u'Trong cách dạy cũ, GV vẽ hình tam giác lên bảng. Trong EduRobot-5, GV cho HS kéo các cạnh để thấy sự thay đổi diện tích.\n'
                          u'* Cách thực hiện cho GV không rành IT: \n'
                          u'- GV mở bài học tương ứng trên web.\n'
                          u'- Gọi HS lên màn hình/tivi thông minh thực hiện thao tác chạm (Touch).\n'
                          u'- Yêu cầu HS khác nhận xét thay đổi ngay trên màn hình.')

    # Biện pháp 2
    doc.add_heading(u'BIỆN PHÁP 2: Bài tập "Kéo - Thả" (Drag and Drop) thay vì chỉ khoanh đáp án', level=2)
    p = doc.add_paragraph(u'Khoanh đáp án A, B, C thường mang tính may rủi. EduRobot-5 sử dụng tính năng kéo-thả số vào công thức.\n'
                          u'* Ý nghĩa: Học sinh phải hiểu công thức nằm ở vị trí nào (Đâu là chiều dài? Đâu là chiều rộng?) thì mới kéo số vào đúng ô được. Nếu sai, hệ thống sẽ đẩy số lại vị trí cũ để HS làm lại.')

    # Biện pháp 3
    doc.add_heading(u'BIỆN PHÁP 3: Củng cố kiến thức bằng Hệ thống Quiz Gamification', level=2)
    p = doc.add_paragraph(u'Thay vì làm phiếu bài tập giấy sau mỗi tiết, EduRobot-5 tích hợp bộ Quiz (trắc nghiệm nhanh) với hiệu ứng âm thanh.\n'
                          u'- Học sinh chọn đúng: Robot nhảy múa và chúc mừng (giúp các em hào hứng).\n'
                          u'- Học sinh chọn sai: Robot đưa ra gợi ý nhỏ bằng ngôn ngữ dễ hiểu để HS tự sửa lỗi.')

    # Biện pháp 4
    doc.add_heading(u'BIỆN PHÁP 4: Chống "Học trước quên sau" bằng Video thực tế', level=2)
    p = doc.add_paragraph(u'Khi dạy về tỉ số phần trăm (giảm giá), EduRobot-5 tích hợp sẵn các đoạn video thực tế trong siêu thị. HS vừa học toán vừa thấy được giá trị của toán học trong cuộc sống, giúp ghi nhớ sâu hơn gấp 3 lần so với chỉ nhìn con số.')

    # --- PHẦN V: QUY TRÌNH 4 BƯỚC ÁP DỤNG TRỰC TIẾP ---
    doc.add_heading(u'V. QUY TRÌNH TỔ CHỨC TIẾT DẠY (Dành cho Giáo viên)', level=1)
    
    doc.add_heading(u'Bước 1: Chuẩn bị nội dung (5 phút)', level=2)
    p = doc.add_paragraph(u'GV truy cập EduRobot-5, chọn bài học theo mã số PPCT (Ví dụ: Tiết 118). GV kiểm tra trước đường truyền mạng và tivi/máy chiếu.')

    doc.add_heading(u'Bước 2: Triển khai bài dạy (Hoạt động chính)', level=2)
    p = doc.add_paragraph(u'- Khởi động: Cho Robot chào cả lớp và tổ chức trò chơi nhỏ 2 phút.\n'
                          u'- Khám phá: GV sử dụng các thanh trượt trên web để HS tự rút ra quy tắc tính.\n'
                          u'- Luyện tập: Cho 5-10 HS lên bảng tương tác trực tiếp với máy tính/tivi.')

    doc.add_heading(u'Bước 3: Đánh giá bằng biểu đồ (Tại lớp)', level=2)
    p = doc.add_paragraph(u'Đây là phần quan trọng nhất. EduRobot-5 sẽ tự động hiển thị: "Có bao nhiêu bạn làm đúng câu 1? Bao nhiêu bạn nhầm lẫn câu 2?". GV dựa vào đây để giảng kỹ lại phần mà nhiều HS sai nhất.')

    doc.add_heading(u'Bước 4: Chia sẻ và Tự học (Sau tiết học)', level=2)
    p = doc.add_paragraph(u'GV gửi mã bài học qua Zalo nhóm lớp. Phụ huynh mở điện thoại cho con chơi lại các trò chơi toán học của tiết vừa qua như một cách ôn tập nhẹ nhàng.')

    # --- PHẦN VI: HIỆU QUẢ THỰC TẾ ---
    doc.add_heading(u'VI. ĐÁNH GIÁ HIỆU QUẢ VÀ SỐ LIỆU MINH CHỨNG', level=1)
    
    # Bảng số liệu
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'Giai đoạn'
    hdr_cells[1].text = u'HT Tốt (%)'
    hdr_cells[2].text = u'Hoàn thành (%)'
    hdr_cells[3].text = u'Chưa HT (%)'

    data = [
        (u'Trước áp dụng', u'17.6%', u'66.4%', u'16.0%'),
        (u'Sau áp dụng', u'28.6%', u'71.4%', u'0%')
    ]
    for stage, ht_tot, ht, cht in data:
        row = table.add_row().cells
        row[0].text = stage
        row[1].text = ht_tot
        row[2].text = ht
        row[3].text = cht

    doc.add_paragraph(u'\nNhìn vào bảng số liệu, ta thấy sự thay đổi thần kỳ: Tỷ lệ học sinh chưa hoàn thành đã giảm về 0. Đặc biệt, năng lực số (khả năng dùng máy tính để học tập) của học sinh tăng lên rõ rệt.')

    # --- PHẦN VII: KẾT LUẬN VÀ KIẾN NGHỊ ---
    doc.add_heading(u'VII. KẾT LUẬN VÀ KIẾN NGHỊ', level=1)
    
    doc.add_heading(u'1. Kết luận', level=2)
    p = doc.add_paragraph(u'EduRobot-5 chính là chiếc cầu nối xóa bỏ khoảng cách giữa "Sách giáo khoa khô khan" và "Bài toán đời thực". Nó giúp giáo viên tiết kiệm thời gian và giúp học sinh yêu thích môn Toán hơn.')

    doc.add_heading(u'2. Kiến nghị', level=2)
    p = doc.add_paragraph(u'- Đối với giáo viên: Không cần quá giỏi IT, hãy cứ mạnh dạn mở web và trải nghiệm cùng học sinh.\n'
                          u'- Đối với nhà trường: Tiếp tục duy trì hạ tầng mạng và khuyến khích giáo viên nhân rộng mô hình này sang các môn khác (như Tiếng Việt, Khoa học).')

    # Lưu file
    file_name = 'CHUYEN_DE_EDUROBOT_CHIE_TIET.docx'
    doc.save(file_name)
    print(f'Done: {os.path.abspath(file_name)}')

if __name__ == "__main__":
    add_detailed_chuyen_de()
