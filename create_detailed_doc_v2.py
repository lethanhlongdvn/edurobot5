# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_font(run, font_name='Times New Roman', size=13):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)

def add_detailed_chuyen_de():
    doc = Document()
    
    # Thiết lập lề trang
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.75)

    # --- TRANG BÌA ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(u'PHÒNG GIÁO DỤC VÀ ĐÀO TẠO HUYỆN CÀNG LONG\nTRƯỜNG TIỂU HỌC ĐỖ VĂN NẠI\n----------***----------\n\n\n')
    set_font(run, size=14)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(u'BÁO CÁO CHUYÊN ĐỀ\n')
    set_font(run, size=18)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(u'ỨNG DỤNG NỀN TẢNG EDUROBOT-5 TRONG DẠY HỌC TOÁN LỚP 5\nTHEO HƯỚNG TƯƠNG TÁC VÀ PHÁT TRIỂN NĂNG LỰC SỐ\n')
    set_font(run, size=20)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(u'\n\n\nNgười thực hiện: Nhóm dự án EduRobot-5\nĐối tượng: Học sinh khối lớp 5\nThời gian: Năm học 2025 - 2026\n')
    set_font(run, size=13)
    
    doc.add_page_break()

    # --- NỘI DUNG CHI TIẾT ---
    # I. Mở đầu
    doc.add_heading(u'I. SỰ CẦN THIẾT VÀ LÝ DO CHỌN CHUYÊN ĐỀ', level=1)
    doc.add_paragraph(u'Môn Toán ở tiểu học đóng vai trò quan trọng trong việc hình thành và phát triển năng lực tư duy toán học cho học sinh...')
    
    # II. Thực trạng
    doc.add_heading(u'II. THỰC TRẠNG TÌNH HÌNH TẠI ĐƠN VỊ', level=1)
    doc.add_paragraph(u'Để diễn giải rõ cho người dùng ít kiến thức IT, chúng ta cần hiểu rằng:')
    doc.add_paragraph(u'1. Phần mềm cũ: Chủ yếu chỉ để trình chiếu ảnh quay video.', style='List Bullet')
    doc.add_paragraph(u'2. EduRobot-5: Cho phép chạm, sửa, thay đổi trực tiếp con số ngay trên màn hình.', style='List Bullet')

    # III. Đặc điểm EduRobot-5
    doc.add_heading(u'III. ĐẶC ĐIỂM HỆ THỐNG EDUROBOT-5', level=1)
    doc.add_paragraph(u'Nhiều người lầm tưởng dạy học công nghệ là phải giỏi lập trình. EduRobot-5 sinh ra để phá vỡ định kiến đó bằng các tính năng:')
    doc.add_paragraph(u'• Robot dẫn đường: Thay giáo viên thực hiện các lời chào, nhắc nhở bài cũ, tạo sự hào hứng cho tiết học ngay từ phút đầu tiên.', style='List Bullet')
    doc.add_paragraph(u'• Lesson Engine: Bộ khung bài giảng đã được lập trình sẵn theo hướng tương tác. Giáo viên chỉ cần mở trình duyệt và bắt đầu dẫn dắt thay vì phải ngồi vẽ slide từng trang.', style='List Bullet')
    doc.add_paragraph(u'• Công cụ kéo-thả: Đây là tinh túy của dự án. Học sinh được cầm con số, kéo vào công thức. Nếu chọn sai chiều dài/chiều rộng, robot sẽ đẩy số ra để học sinh suy nghĩ lại.', style='List Bullet')

    # IV. Các biện pháp chi tiết
    doc.add_heading(u'IV. CÁC BIỆN PHÁP NÂNG CAO CHẤT LƯỢNG DẠY HỌC', level=1)
    
    doc.add_heading(u'Biện pháp 1: Chuyển đổi từ "Hình ảnh tĩnh" sang "Mô phỏng động"', level=2)
    doc.add_paragraph(u'Phân ích cho người không chuyên: Hình ảnh tĩnh là bức tranh chết trên trang sách. Mô phỏng động là khi bạn thay đổi một cạnh của hình thang, diện tích của nó sẽ nhảy số ngay trước mắt bạn. HS sẽ hiểu "À, cạnh này dài hơn thì diện tích sẽ lớn hơn" một cách rất tự nhiên.')
    
    doc.add_heading(u'Biện pháp 2: Sử dụng bộ Quiz (Trò chơi) tích hợp âm thanh', level=2)
    doc.add_paragraph(u'Sử dụng âm thanh (Tiếng vỗ tay, tiếng robot cười) để kích thích hệ thần kinh hưng phấn của trẻ em. Khi các em học trong trạng thái vui vẻ, kiến thức sẽ được tiếp nhận nhanh gấp nhiều lần.')

    doc.add_heading(u'V. QUY TRÌNH ÁP DỤNG TRONG TIẾT DẠY', level=1)
    doc.add_paragraph(u'- Bước 1: Mở trình duyệt web và đăng nhập vào EduRobot-5 (Dễ như mở Facebook).')
    doc.add_paragraph(u'- Bước 2: Sử dụng các phím điều hướng để đi qua 4 hoạt động của bài học: Khởi động - Khám phá - Luyện tập - Vận dụng.')
    doc.add_paragraph(u'- Bước 3: Cho học sinh lên bảng tương tác với tivi/máy tính.')

    # VI. Kết luận
    doc.add_heading(u'VI. ĐÁNH GIÁ HIỆU QUẢ', level=1)
    doc.add_paragraph(u'Sự thay đổi không chỉ nằm ở điểm số, mà ở năng lực số của trẻ em vùng sâu, vùng xa Càng Long được thu hẹp khoảng cách với thành phố nhờ vào EduRobot-5.')

    # VII. Kiến nghị
    doc.add_heading(u'VII. KIẾN NGHỊ', level=1)
    doc.add_paragraph(u'Kiến nghị nhà trường tiếp tục ủng hộ đội ngũ phát triển dự án để EduRobot-5 trở thành công cụ dạy học số 1 trong cụm chuyên môn.')

    doc.save('CHUYEN_DE_EDUROBOT_5_SIE_CHI_TIET.docx')
    print('OK: CHUYEN_DE_EDUROBOT_5_SIE_CHI_TIET.docx')

if __name__ == "__main__":
    add_detailed_chuyen_de()
